#!/usr/bin/env python3
"""
DOCX (Word Document) Converter
Handles Microsoft Word .docx files
"""

from typing import Tuple, Dict, Any, TYPE_CHECKING

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Only import for type checking to avoid errors when python-docx not installed
if TYPE_CHECKING:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

from .base import BaseConverter
from ..core import (
    normalize_text_for_llm,
    detect_language,
    apply_bidi_algorithm,
    generate_metadata,
    create_page_separator
)


class DOCXConverter(BaseConverter):
    """
    DOCX (Microsoft Word) document converter
    """

    supported_extensions = ['.docx']

    def __init__(self, apply_bidi: bool = False, normalize: bool = True):
        """
        Initialize DOCX converter

        Args:
            apply_bidi: Apply BiDi algorithm for RTL text (default: False)
                       DOCX files store text in logical order, which is what
                       markdown expects. Only enable if converting from sources
                       that use visual order (e.g., some PDFs)
            normalize: Normalize text (remove niqqud, excess whitespace, etc.)
        """
        super().__init__(apply_bidi=apply_bidi, normalize=normalize)

        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX conversion. "
                "Install it with: pip install python-docx"
            )

    def _extract_paragraph_text(self, paragraph: "Paragraph") -> str:
        """
        Extract text from a paragraph

        Args:
            paragraph: Paragraph object from python-docx

        Returns:
            Extracted text string
        """
        return paragraph.text

    def _extract_table_text(self, table: "Table") -> str:
        """
        Extract text from a table and format as markdown table

        Args:
            table: Table object from python-docx

        Returns:
            Markdown-formatted table string
        """
        if not table.rows:
            return ""

        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))

        if len(rows_text) > 1:
            # Add markdown table header separator
            header = rows_text[0]
            separator = " | ".join(["---"] * len(table.rows[0].cells))
            body = rows_text[1:]
            return f"\n| {header} |\n| {separator} |\n" + "\n".join(f"| {row} |" for row in body) + "\n"
        else:
            return "\n| " + rows_text[0] + " |\n"

    def _extract_document_structure(self, doc: "Document") -> Tuple[str, int]:
        """
        Extract text from document preserving structure

        Args:
            doc: Document object from python-docx

        Returns:
            Tuple of (extracted_text, section_count)
        """
        all_text = ""
        section_count = 0
        paragraph_count = 0

        print(f"Extracting text from DOCX document...")

        # Iterate through document elements
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                paragraph = doc.paragraphs[paragraph_count]
                text = self._extract_paragraph_text(paragraph)

                if text and text.strip():
                    # Check if it's a heading based on style
                    if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                        # Extract heading level
                        level = paragraph.style.name.replace('Heading ', '').strip()
                        if level.isdigit():
                            all_text += f"\n{'#' * int(level)} {text}\n\n"
                        else:
                            all_text += f"\n## {text}\n\n"

                        section_count += 1
                    else:
                        all_text += text + "\n\n"

                paragraph_count += 1

            # Handle tables
            elif element.tag.endswith('tbl'):
                # Find the corresponding table
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        all_text += table_text + "\n"
                break  # Only process first table match

        print(f"  Extracted {paragraph_count} paragraphs")
        print(f"  Found {len(doc.tables)} tables")
        print(f"  Identified {section_count} sections")

        return all_text, max(1, section_count)

    def extract(self, file_path: str, **kwargs) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from DOCX file

        Args:
            file_path: Path to DOCX file
            **kwargs: Additional parameters (ignored for compatibility)

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        self.validate_file(file_path)

        # Open and parse DOCX file
        doc = Document(file_path)

        # Extract text preserving structure
        all_text, section_count = self._extract_document_structure(doc)

        # Process text
        if self.normalize:
            all_text = normalize_text_for_llm(all_text)

        if self.apply_bidi:
            # Note: DOCX files store text in logical order (storage order).
            # BiDi algorithm converts logical→visual for display purposes.
            # Since markdown also expects logical order, BiDi should only be
            # applied in special cases (e.g., when text source uses visual order)
            all_text = apply_bidi_algorithm(all_text)

        # Detect language
        primary_lang, detected_langs = detect_language(all_text[:5000])

        # Generate metadata
        metadata = generate_metadata(
            source_file=file_path,
            pages=section_count,  # Use section count as "pages" equivalent
            language=primary_lang,
            detected_languages=detected_langs,
            method="docx_extraction",
            characters=len(all_text),
            tables=len(doc.tables),
            paragraphs=len(doc.paragraphs)
        )

        return all_text, metadata
